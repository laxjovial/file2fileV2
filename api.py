from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
import os
import pandas as pd
from docx import Document
from pdf2docx import Converter
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import pypandoc
import uvicorn

# Initialize FastAPI app
app = FastAPI(title="File2File Conversion API")

# Supported formats (must match frontend)
doc_types = ["pdf", "docx", "txt"]
sheet_types = ["csv", "xls", "xlsx"]

# Function for document (PDF, DOCX, TXT) conversions - Reused from Streamlit app
def convert_doc_file_backend(file_bytes_io: BytesIO, source: str, target: str, font_size: int = 12):
    result = BytesIO()
    file_bytes_io.seek(0) # Ensure pointer is at the beginning

    unique_id = os.urandom(8).hex()
    temp_input_path = f"temp_input_{unique_id}.{source}"
    temp_output_path = f"temp_output_{unique_id}.{target}"

    try:
        if source == "pdf":
            with open(temp_input_path, "wb") as f:
                f.write(file_bytes_io.read())

            if target == "docx":
                cv = Converter(temp_input_path)
                cv.convert(temp_output_path, start=0, end=None)
                cv.close()
                with open(temp_output_path, "rb") as f:
                    result.write(f.read())
            elif target == "txt":
                with pdfplumber.open(file_bytes_io) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                    result.write(text.encode("utf-8"))

        elif source == "docx":
            with open(temp_input_path, "wb") as f:
                f.write(file_bytes_io.read())

            if target == "pdf":
                pypandoc.convert_file(
                    temp_input_path,
                    "pdf",
                    outputfile=temp_output_path,
                    extra_args=['--pdf-engine=wkhtmltopdf']
                )
                with open(temp_output_path, "rb") as f:
                    result.write(f.read())
            elif target == "txt":
                doc_bytes = BytesIO(file_bytes_io.read())
                doc = Document(doc_bytes)
                text = "\n".join([p.text for p in doc.paragraphs])
                result.write(text.encode("utf-8"))

        elif source == "txt":
            text = file_bytes_io.read().decode("utf-8")
            if target == "pdf":
                c = canvas.Canvas(result, pagesize=letter)
                c.setFont("Helvetica", font_size) # Apply font size from frontend
                y = 750
                for line in text.splitlines():
                    c.drawString(50, y, line)
                    y -= (font_size + 5) # Adjust line spacing based on font size
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", font_size)
                        y = 750
                c.save()
            elif target == "docx":
                doc = Document()
                # Apply basic styling based on Markdown-like syntax if present
                for line in text.splitlines():
                    if line.startswith('**') and line.endswith('**'):
                        paragraph = doc.add_paragraph(line.strip('**'))
                        paragraph.runs[0].bold = True
                    elif line.startswith('*') and line.endswith('*'):
                        paragraph = doc.add_paragraph(line.strip('*'))
                        paragraph.runs[0].italic = True
                    else:
                        doc.add_paragraph(line)
                doc.save(result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")
    finally:
        if os.path.exists(temp_input_path):
            os.remove(temp_input_path)
        if os.path.exists(temp_output_path):
            os.remove(temp_output_path)

    result.seek(0)
    return result

# Function for spreadsheet (CSV, XLS, XLSX) conversions - Reused from Streamlit app
def convert_sheet_file_backend(file_bytes_io: BytesIO, source: str, target: str):
    result = BytesIO()
    file_bytes_io.seek(0)

    if source == "csv":
        df = pd.read_csv(file_bytes_io)
    else:
        df = pd.read_excel(file_bytes_io)

    if target == "csv":
        df.to_csv(result, index=False)
    else:
        df.to_excel(result, index=False, engine="openpyxl")

    result.seek(0)
    return result

@app.post("/convert")
async def convert_file_endpoint(
    file: UploadFile = File(...),
    from_format: str = Form(...),
    to_format: str = Form(...),
    font_size: int = Form(12) # Receive font size from frontend
):
    # Read the uploaded file content into BytesIO
    file_content = await file.read()
    file_bytes_io = BytesIO(file_content)

    if from_format in doc_types and to_format in doc_types:
        converted_output = convert_doc_file_backend(file_bytes_io, from_format, to_format, font_size)
        media_type = f"application/{to_format}" if to_format != "txt" else "text/plain"
    elif from_format in sheet_types and to_format in sheet_types:
        converted_output = convert_sheet_file_backend(file_bytes_io, from_format, to_format)
        media_type = f"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" if to_format == "xlsx" else f"text/{to_format}"
    else:
        raise HTTPException(status_code=400, detail="Cross-type conversions (e.g., DOCX to CSV) are not supported.")

    return StreamingResponse(converted_output, media_type=media_type, 
                             headers={"Content-Disposition": f"attachment; filename=converted.{to_format}"})

# To run this FastAPI app locally:
# Save this file as api.py
# Open your terminal in the same directory and run:
# uvicorn api:app --host 0.0.0.0 --port 8000
# This will start the API server on http://localhost:8000
